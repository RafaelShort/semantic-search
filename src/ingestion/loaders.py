"""
Loaders de documentos.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import httpx
from bs4 import BeautifulSoup
from docx import Document
from loguru import logger
from pypdf import PdfReader

# Estrutura de dados padrão retornada por todos os loaders

@dataclass
class DocumentData:
    """
    Representa um documento carregado, independente da sua origem.

    Todos os loaders retornam este mesmo formato,
    garantindo que o pipeline funcione de forma uniforme.
    """
    source: str             
    content: str             
    doc_type: str            
    metadata: dict = field(default_factory=dict)

    def is_valid(self) -> bool:
        """Verifica se o documento tem conteúdo útil."""
        return bool(self.content and len(self.content.strip()) > 50)

# Utilitários de limpeza de texto

def clean_text(text: str) -> str:
    """
    Limpa e normaliza texto extraído de documentos.

    Operações:
    - Remove caracteres de controle
    - Normaliza espaços em branco múltiplos
    - Normaliza quebras de linha
    - Remove linhas vazias excessivas
    """
    if not text:
        return ""

    # Remove caracteres de controle (exceto \n e \t)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # Normaliza espaços múltiplos em um único espaço
    text = re.sub(r" +", " ", text)

    # Normaliza tabulações para espaço
    text = re.sub(r"\t+", " ", text)

    # Remove mais de 2 quebras de linha consecutivas
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Remove espaços no início/fim de cada linha
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)

    return text.strip()

# Loaders

class TextLoader:
    """Carrega arquivos .txt"""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml"}

    def load(self, file_path: str | Path) -> DocumentData:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")

        logger.info(f"📄 Carregando TXT: {path.name}")

        # Tenta diferentes encodings
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
        content = None

        for encoding in encodings:
            try:
                content = path.read_text(encoding=encoding)
                logger.debug(f"   Encoding detectado: {encoding}")
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            raise ValueError(f"Não foi possível decodificar: {file_path}")

        return DocumentData(
            source=str(path.absolute()),
            content=clean_text(content),
            doc_type="txt",
            metadata={
                "filename":  path.name,
                "extension": path.suffix,
                "size_bytes": path.stat().st_size,
            }
        )


class PDFLoader:
    """Carrega arquivos .pdf"""

    def load(self, file_path: str | Path) -> DocumentData:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")

        logger.info(f"Carregando PDF: {path.name}")

        reader = PdfReader(str(path))
        pages_text = []

        for page_num, page in enumerate(reader.pages):
            try:
                page_content = page.extract_text()
                if page_content and page_content.strip():
                    # Adiciona separador entre páginas
                    pages_text.append(
                        f"[Página {page_num + 1}]\n{page_content}"
                    )
            except Exception as exc:
                logger.warning(f"   Erro na página {page_num + 1}: {exc}")
                continue

        full_content = "\n\n".join(pages_text)

        # Extrai metadados do PDF
        metadata = {
            "filename":   path.name,
            "pages":      len(reader.pages),
            "size_bytes": path.stat().st_size,
        }

        if reader.metadata:
            if reader.metadata.title:
                metadata["title"] = reader.metadata.title
            if reader.metadata.author:
                metadata["author"] = reader.metadata.author

        logger.info(f"   ✅ {len(reader.pages)} páginas extraídas")

        return DocumentData(
            source=str(path.absolute()),
            content=clean_text(full_content),
            doc_type="pdf",
            metadata=metadata
        )


class DocxLoader:
    """Carrega arquivos .docx (Microsoft Word)"""

    def load(self, file_path: str | Path) -> DocumentData:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")

        logger.info(f"📘 Carregando DOCX: {path.name}")

        doc = Document(str(path))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserva títulos com marcação especial
                if para.style.name.startswith("Heading"):
                    paragraphs.append(f"\n## {text}\n")
                else:
                    paragraphs.append(text)

        full_content = "\n".join(paragraphs)

        return DocumentData(
            source=str(path.absolute()),
            content=clean_text(full_content),
            doc_type="docx",
            metadata={
                "filename":        path.name,
                "paragraphs":      len(doc.paragraphs),
                "size_bytes":      path.stat().st_size,
            }
        )


class URLLoader:
    """Carrega conteúdo de páginas web via URL"""

    def __init__(self, timeout: int = 15):
        self.timeout = timeout

    def load(self, url: str) -> DocumentData:
        logger.info(f"Carregando URL: {url}")

        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36"
            )
        }

        try:
            response = httpx.get(
                url,
                headers=headers,
                timeout=self.timeout,
                follow_redirects=True
            )
            response.raise_for_status()
        except httpx.TimeoutException:
            raise TimeoutError(f"Timeout ao acessar URL: {url}")
        except httpx.HTTPStatusError as exc:
            raise ValueError(f"Erro HTTP {exc.response.status_code}: {url}")

        # Parseia o HTML e extrai apenas o texto
        soup = BeautifulSoup(response.text, "html.parser")

        # Remove elementos que não têm conteúdo útil
        for tag in soup(["script", "style", "nav", "footer",
                          "header", "aside", "form", "iframe"]):
            tag.decompose()

        # Tenta extrair o título da página
        title = ""
        if soup.title and soup.title.string:
            title = soup.title.string.strip()

        # Extrai texto principal
        main_content = (
            soup.find("main") or
            soup.find("article") or
            soup.find(id="content") or
            soup.find(class_="content") or
            soup.body
        )

        if main_content:
            text = main_content.get_text(separator="\n")
        else:
            text = soup.get_text(separator="\n")

        logger.info(f"   Página carregada | Título: '{title}'")

        return DocumentData(
            source=url,
            content=clean_text(text),
            doc_type="url",
            metadata={
                "url":         url,
                "title":       title,
                "status_code": response.status_code,
                "content_type": response.headers.get("content-type", ""),
            }
        )

# Factory — escolhe o loader certo automaticamente

class DocumentLoader:
    """
    Facade que escolhe o loader correto baseado no tipo de entrada.

    Uso:
        loader = DocumentLoader()
        doc = loader.load("arquivo.pdf")
        doc = loader.load("https://exemplo.com")
        doc = loader.load("texto.txt")
    """

    def __init__(self):
        self._text_loader = TextLoader()
        self._pdf_loader  = PDFLoader()
        self._docx_loader = DocxLoader()
        self._url_loader  = URLLoader()

    def load(self, source: str) -> Optional[DocumentData]:
        """
        Carrega um documento de qualquer fonte suportada.

        Args:
            source: Caminho de arquivo ou URL

        Returns:
            DocumentData ou None se falhar
        """
        try:
            # Verifica se é URL
            if source.startswith(("http://", "https://")):
                return self._url_loader.load(source)

            # Se é um arquivo, escolhe loader pela extensão
            path = Path(source)
            extension = path.suffix.lower()

            if extension in TextLoader.SUPPORTED_EXTENSIONS:
                return self._text_loader.load(path)
            elif extension == ".pdf":
                return self._pdf_loader.load(path)
            elif extension in {".docx", ".doc"}:
                return self._docx_loader.load(path)
            else:
                logger.warning(f"⚠️  Tipo não suportado: '{extension}' | {source}")
                return None

        except FileNotFoundError as exc:
            logger.error(f"Arquivo não encontrado: {exc}")
            return None
        except Exception as exc:
            logger.error(f"Erro ao carregar '{source}': {exc}")
            return None

    def load_directory(
        self,
        directory: str | Path,
        recursive: bool = False
    ) -> list[DocumentData]:
        """
        Carrega todos os arquivos suportados de uma pasta.

        Args:
            directory:  Caminho da pasta
            recursive:  Se True, busca em subpastas também
        """
        dir_path = Path(directory)
        if not dir_path.exists():
            raise FileNotFoundError(f"Pasta não encontrada: {directory}")

        # Extensões suportadas
        extensions = (
            TextLoader.SUPPORTED_EXTENSIONS |
            {".pdf", ".docx", ".doc"}
        )

        # Busca arquivos
        pattern = "**/*" if recursive else "*"
        files = [
            f for f in dir_path.glob(pattern)
            if f.is_file() and f.suffix.lower() in extensions
        ]

        logger.info(f"{len(files)} arquivos encontrados em '{dir_path.name}'")

        documents = []
        for file in files:
            doc = self.load(str(file))
            if doc and doc.is_valid():
                documents.append(doc)
            elif doc:
                logger.warning(f"Documento muito curto, ignorado: {file.name}")

        logger.info(f"{len(documents)}/{len(files)} documentos carregados")
        return documents
